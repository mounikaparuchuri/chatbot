import streamlit as st
from google import genai
from google.genai import types
from google.genai.types import (
    CreateBatchJobConfig,
    CreateCachedContentConfig,
    EmbedContentConfig,
    FunctionDeclaration,
    GenerateContentConfig,
    HarmBlockThreshold,
    HarmCategory,
    Part,
    SafetySetting,
    Tool,
)
from PIL import Image
import PyPDF2
import docx
from io import BytesIO
from read import setup_db, save_data, retrieve_data, delete_chat_log

# Global variables
global username, db_file_name
db_file_name = 'chat_history.db'
username = ''

def setup_app(prompt_input=None):    
    client = genai.Client(
        api_key=st.secrets["GEMINI_API_KEY"],
    )
    # Load system prompt from secrets
    if "pname" in query_params:
        system_prompt_name = query_params["pname"]
        st.session_state.system_prompt = st.secrets.get(system_prompt_name)

    if "username" in query_params and "clearall" in query_params:
        username = query_params["username"]
        db_file_name = f"{username}_chat_history.db"
        delete_chat_log(username, db_file_name)


    # Initialize messages and load history if it's the first run
    if "messages" not in st.session_state:
        st.session_state.messages = retrieve_data(db_file_name) or []

    generate_content_config = types.GenerateContentConfig(
        system_instruction=st.session_state.system_prompt,
        temperature = 1,
        top_p = 0.95,
        max_output_tokens = 8192,
        safety_settings = [types.SafetySetting(
        category="HARM_CATEGORY_HATE_SPEECH",
        threshold="OFF"
        ),types.SafetySetting(
        category="HARM_CATEGORY_DANGEROUS_CONTENT",
        threshold="OFF"
        ),types.SafetySetting(
        category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
        threshold="OFF"
        ),types.SafetySetting(
        category="HARM_CATEGORY_HARASSMENT",
        threshold="OFF"
        )],
    )
    model = "gemini-2.0-flash-lite-001"
    contents = [
        types.Content(
        role="user",
        parts=[{"text": ""}]
        )
    ]
    chat = client.chats.create(model = model, config = generate_content_config,)
    response = chat.send_message(prompt_input.text)

    with st.chat_message("user"):
        st.markdown(prompt_input.text)
    print(response.text)

        # Add the text from the prompt.
    user_content.append(prompt_input.text)


        # Save the new user message and the assistant's response
    if st.session_state.messages:
        save_data(db_file_name, username, st.session_state.messages, response.text)
        st.session_state.messages.append({"role": "user", "content": st.session_state.messages})
        st.session_state.messages.append({"role": "assistant", "content": response.text})
    else: # This branch handles the initial response from the system prompt
        st.session_state.messages.append({"role": "assistant", "content": response.text})

    display_messages()


def display_messages():
    """Displays chat messages from session state, excluding the system prompt."""
    for message in st.session_state.messages:
        if message["role"] == "user":
            with st.chat_message("user"):
                content = message["content"]
                if isinstance(content, list):
                    text_content = next((part for part in content if isinstance(part, str)), None)
                    if text_content:
                        st.markdown(text_content)
                else:
                    st.markdown(content)
        elif message["role"] == "assistant":
            with st.chat_message("assistant"):
                st.markdown(message["content"])

def handle_file_uploads(uploaded_files):
    """Processes uploaded files and returns a list of content parts."""
    user_content = []
    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_extension = uploaded_file.name.split(".")[-1].lower()
            
            if file_extension in ["jpg", "jpeg", "png"]:
                image = Image.open(uploaded_file)
                user_content.append(image)
            elif file_extension == "pdf":
                try:
                    pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.getvalue()))
                    pdf_text = "".join(p.extract_text() or "" for p in pdf_reader.pages)
                    user_content.append(f"Content from PDF: {pdf_text}")
                except Exception as e:
                    st.error(f"Error reading PDF: {e}")
                    user_content.append(f"Could not read PDF. Error: {e}")
            elif file_extension in ["doc", "docx"]:
                try:
                    doc = docx.Document(uploaded_file)
                    doc_text = " ".join(p.text for p in doc.paragraphs)
                    user_content.append(f"Content from Word document: {doc_text}")
                except Exception as e:
                    st.error(f"Error reading Word document: {e}")
                    user_content.append(f"Could not read Word document. Error: {e}")
    return user_content


"""Initializes the app by checking for a username, setting up the DB, and loading chat history."""
query_params = st.query_params
st.session_state.system_prompt = ""

if "username" in query_params and "pname" in query_params:
    username = query_params["username"]
    db_file_name = f"{username}_chat_history.db"
    setup_db(db_file_name)
    st.title("💬 Chatbot")
    st.write("Find My Chapter 3 chatbot.")
else:
    st.title("💬 Chatbot")
    st.write("Error occurred: parameters are missing.")
    st.stop()

prompt_input = st.chat_input(
    "Say something and/or attach an image",
    accept_file=True,
    file_type=["jpg", "jpeg", "png", "pdf", "doc", "docx"],
)

if prompt_input and prompt_input.text:
    # Create the user message content.
    user_content = []

    # Check for uploaded files
    if prompt_input["files"]:
        uploaded_file = prompt_input["files"][0]
        file_extension = uploaded_file.name.split(".")[-1].lower()

        # Handle images
        if file_extension in ["jpg", "jpeg", "png"]:
            image = Image.open(uploaded_file)
            user_content.append(image)
            with st.chat_message("user"):
                st.image(image)

        # Handle PDFs
        elif file_extension == "pdf":
            pdf_text = ""
            try:
                # We need to use BytesIO to read the file in-memory
                pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.getvalue()))
                for page in pdf_reader.pages:
                    pdf_text += page.extract_text() or ""
                user_content.append(f"Content from PDF: {pdf_text}")
                with st.chat_message("user"):
                    st.success("PDF processed successfully.")
            except Exception as e:
                st.error(f"Error reading PDF: {e}")
                user_content.append(f"Could not read PDF. Error: {e}")
                
        # Handle Word documents
        elif file_extension in ["doc", "docx"]:
            doc_text = ""
            try:
                # Streamlit's UploadedFile object is file-like.
                # docx library can read from it directly.
                doc = docx.Document(uploaded_file)
                for paragraph in doc.paragraphs:
                    doc_text += paragraph.text + " "
                user_content.append(f"Content from Word document: {doc_text}")
                with st.chat_message("user"):
                    st.success("Word document processed successfully.")
            except Exception as e:
                st.error(f"Error reading Word document: {e}")
                user_content.append(f"Could not read Word document. Error: {e}")

    # Main app logic
    setup_app(prompt_input)
    display_messages()
    
    # Get LLM response
    # get_llm_response(user_input_content=user_content_list)
    # st.rerun()
